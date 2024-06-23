import dropbox
import fitz  # PyMuPDF
from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import UnstructuredFileLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import CacheBackedEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma, FAISS
from langchain.storage import LocalFileStore
from langchain.chains import RetrievalQA
from langchain.prompts import ChatPromptTemplate
from langchain.schema.runnable import RunnablePassthrough
from langchain.prompts.chat import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.memory import ConversationBufferMemory
from langchain.agents import create_openai_functions_agent, AgentExecutor
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder

import json
import os
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.pydantic_v1 import BaseModel, Field
from typing import List
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.shared import Inches
from langchain.tools.retriever import create_retriever_tool
from langchain_community.agent_toolkits.load_tools import load_tools
from langchain import hub
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
import io
from dropbox import Dropbox
from dropbox.files import WriteMode
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import openai
from dotenv import load_dotenv
load_dotenv()

openai_api_key = os.environ.get('OPENAI_API_KEY')
openai_api_key2 = os.environ.get('OPENAI_API_KEY2')

# Set your Dropbox access token and folder path
ACCESS_TOKEN = "EpsonAPI"

def MakeStudentInfo(Question_pdf, Answer_json):
    #document = fitz.open(Question_pdf)
    document = Question_pdf
    # 텍스트를 저장할 파일 열기
    Question_pdf_file = "Epson-data/Question_pdf.txt"
    with open(Question_pdf_file, 'w', encoding='utf-8') as file:
        # PDF 파일의 각 페이지에서 텍스트 추출
        for page_number in range(len(document)):
            page = document.load_page(page_number)
            text = page.get_text()
            file.write(f"Page {page_number + 1}:\n")
            file.write(text)
            file.write("\n\n")
    
    document.close()
    with open(Question_pdf_file, 'r', encoding='utf-8') as file:
        extracted_text = file.read()

    # JSON 파일 읽기
    #with open(Answer_json, 'r', encoding='utf-8') as file:
    #    json_data = json.load(file)

    # Pydantic 모델 정의
    class StudentResponse(BaseModel):
        Number: int = Field(description="문제 번호")
        Question: str = Field(description="문제")
        Text: str = Field(description="문제에 첨부된 지문")
        Candidate: List[str] = Field(description="5가지 정답 후보")
        StudentAnswer: int = Field(description="학생이 제출한 정답")

    llm = ChatOpenAI(openai_api_key=openai_api_key, 
                    temperature=0.1, 
                    model="gpt-4o")
    
    human_message_prompt = """
    {extracted_text} 이것은 N개의 문제이며, 
    {Answer_sheet_json} 이것은 N개의 문제에 대해 제출한 학생의 답안지이다. 
    {Answer_sheet_json} 위 파일을 보면, 몇개의 문제인지 인식 가능하다. 
    N개의 문제를 인식해서 문제 번호, 문제, 보기, 학생이 제출한 정답을 모두 출력해줘.
    Generate a JSON object with the following fields:
    - Number: 각 문제 번호
    - Question: 각 문제
    - Text: 각 문제에 첨부된 지문
    - Candidate: 각 문제에 대한 보기
    - StudentAnswer: 학생이 제출한 정답

    Here is a Example:
    - Number: 1
    - Question: '<보기>를 바탕으로 접사에 대해 탐구한 내용으로 적절하지 않은 것은?'
    - Text: '<보기>\n접두사와 접미사는 다양한 품사의 어근에 결합하여 파생어를 형성한다. 접두사는 어근에 어휘 적 의미를 더해 주는 경우가 많으며, 파생어의 품사를 결정하는 경우는 거의 없다. 한편 접미사는 파생 과정에서 어근에 문법적 의미를 더해 주는 경우가 있으며, 파생어의 품사가 어근의 품사와 달라지는 경우가 있다.'
    - Candidate: ["① '덧저고리, 덧대다'를 보니 접두사 '덧-'은 명사에도 결합하고 동사에도 결합하는구나.", 
                  "② '군말, 군살'을 보니 접두사 '군-'은 '쓸데없는'이라는 어휘적 의미를 어근에 더해 주는구나.", 
                  "③ '잠꾸러기, 욕심꾸러기'를 보니 어근에 접미사 '-꾸러기'가 결합하면 파생어의 품사가 달라 지는구나.", 
                  "④ '늦추다, 낮추다'를 보니 접미사 '-추-'는 어근 '늦-'과 '낮-'에 사동이라는 문법적 의미를 더해 주는구나.", 
                  "⑤ '풋과일, 헛디디다'를 보니 접두사 '풋-'과 '헛-'이 어근에 결합해도 파생어의 품사는 어근의 품사와 달라지지 않는구나."]
    - StudentAnswer: 3"""

    parser = JsonOutputParser(pydantic_object=StudentResponse)
    #format_instructions = parser.get_format_instructions()
    chat_prompt_template = ChatPromptTemplate.from_messages([human_message_prompt])
    chain = chat_prompt_template | llm | parser
    response = chain.invoke({"extracted_text": extracted_text, 
                             "Answer_sheet_json": Answer_json})
    # JSON 파일로 저장
    with open('./QuestionAnswer.json', 'w', encoding='utf-8') as json_file:
        json.dump(response, json_file, ensure_ascii=False, indent=4)
    
    print(1)
    return response

def MakeScoreCommentary(QuestionAnswerJson):

    # JSON 파일 열기
    #with open(QuestionAnswerJson, 'r') as file:
        # JSON 파일 읽기 및 파싱
        #data = json.load(file)
    data = QuestionAnswerJson
        # Pydantic 모델 정의
    class CommentaryResponse(BaseModel):
        Number: int = Field(description="문제 번호")
        StudentAnswer: int = Field(description="학생이 제출한 정답")
        CorrectAnswer: int = Field(description="실제 정답")
        IsCorrect: bool = Field(description="정답 여부")
        CommentarySummarize: List[str] = Field(description="문제 해설 3줄 요약")
    
    cache_dir = LocalFileStore("./.cache/")

    splitter = CharacterTextSplitter.from_tiktoken_encoder(
        separator="\n",
        chunk_size=600,
        chunk_overlap=100,
    )
    loader = UnstructuredFileLoader("data/2024 수능특강 언어와 매체 정답.pdf")

    docs = loader.load_and_split(text_splitter=splitter)

    embeddings = OpenAIEmbeddings()

    cached_embeddings = CacheBackedEmbeddings.from_bytes_store(embeddings, cache_dir)

    vectorstore = FAISS.from_documents(docs, cached_embeddings)

    retriver = vectorstore.as_retriever()
    
    
    llm = ChatOpenAI(openai_api_key=openai_api_key2, 
                 model= "gpt-4o", 
                 temperature=0.1)
    structured_llm = llm.with_structured_output(CommentaryResponse)
    # parser가 get_format_instructions를 통해 스키마 생성, LLM 출력을 Json 형식으로 파싱
    parser = JsonOutputParser(pydantic_object=CommentaryResponse)
    format_instructions = parser.get_format_instructions()
    prompt = PromptTemplate(
    template = """Answer the question based only on the following context:
    {context}

    Question: {question}\n{format_instructions}
    """,
        input_variables=["context", "question"],
        partial_variables={"format_instructions": format_instructions},
    )
    chain = (
        {
            "context": retriver,
            "question": RunnablePassthrough(),
        }
        | prompt
        | structured_llm
        #| parser
    )
    promptforAnswer = """
    위 내용은 문제, 보기, 학생이 제출한 답안이다. 
    정답지에서 ① = 1, ② = 2, ③ = 3, ④ = 4, ⑤ = 5 을 의미한다.
    정답지를 활용해 문제 번호, 학생이 제출한 답안, ,실제 정답, 학생의 정답 여부을 작성해줘.
    문제해설을 작성해주되, 학생의 답안은 정확하다. 와 같은 해설은 보여주지 말아줘. 3줄로 요약해줘.
    아래 예시처럼 작성해줘.
    - Number: 1
    - StudentAnswer: 2
    - CorrectAnswer: 5
    - IsCorrect: False
    - CommentarySummarize: 
    ["토론의 내용은 아동의 개인 정보 노출과 관련된 법적 보호 필요성에 초점을 맞추고 있음.",
    "찬성 측은 개인 정보 보호의 중요성을 강조하며 법적 보호를 주장할 수 있음.",
    "반대 측은 개인 정보의 미숙한 권리 행사나 삭제 요구 권리의 제도화를 주장할 수 있음."]
    """

    Question_list = []
    
    # 각 객체의 'name' 값을 리스트에 추가
    
    for item in data:
        Question = f"문제 번호: {item['Number']} \n문제: {item['Question']} \n지문: {item['Text']} \n보기: {item['Candidate']} \n학생이 제출한 답안: {item['StudentAnswer']}\n" + promptforAnswer
        Question_list.append(Question)

    
    response = chain.batch(Question_list)
    
    #print(Total_response)
    response_list = []
    for resp in response:
        resp_dict = {"Number": resp.Number,
                     "StudentAnswer": resp.StudentAnswer,
                     "CorrectAnswer": resp.CorrectAnswer,
                     "IsCorrect": resp.IsCorrect,
                     "CommentarySummarize": resp.CommentarySummarize}
        response_list.append(resp_dict)
    
    # JSON 파일로 저장
    with open('./AnswerCommentary.json', 'w', encoding='utf-8') as json_file:
        json.dump(response_list, json_file, ensure_ascii=False, indent=4)
    
    print(2)
    return response_list


def CreateMemorizationBookAddCommentary(QuestionAnswer_json, AnswerCommentary_json, Output_path):
    # JSON 파일 열기
    with open(QuestionAnswer_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        QuestionAnswer_json_data = json.load(file)

    with open(AnswerCommentary_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        AnswerCommentary_json_data = json.load(file)

    doc = Document()
    dbx = dropbox.Dropbox(ACCESS_TOKEN)
    for i in range(len(QuestionAnswer_json_data)):
        if AnswerCommentary_json_data[i]['IsCorrect'] == False:
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(5)

            # 날짜, 월, 일
            now = datetime.now()
            current_date = now.date()
            cell = table.cell(0, 0)
            cell.text = "날짜"
            cell = table.cell(0, 1)
            cell.text = f"{current_date.year}년 {current_date.month}월 {current_date.day}일"
            
            # 과목
            cell = table.cell(1, 0)
            cell.text = "과목"
            cell = table.cell(1, 1)
            cell.text = "언어와 매체"

            cell = table.cell(2, 0)
            cell.text = "문제"
            cell = table.cell(2, 1)
            # 문제 + 지문 + 보기 5개 + 정답 + 해설
            cell.text = QuestionAnswer_json_data[i]['Question'] + "\n\n" + QuestionAnswer_json_data[i]['Text']

            cell = table.cell(3, 0)
            cell.text = "보기"
            cell = table.cell(3, 1)
            cell.text = "\n".join(QuestionAnswer_json_data[i]['Candidate'])
            # 난이도
            cell = table.cell(4, 0)
            cell.text = "정답"
            cell = table.cell(4, 1)
            cell.text = str(AnswerCommentary_json_data[i]['CorrectAnswer'])
            cell = table.cell(5, 0)
            cell.text = "해설"
            cell = table.cell(5, 1)
            cell.text = "\n".join(AnswerCommentary_json_data[i]['CommentarySummarize'])

            doc.add_page_break()
    
    #doc.save(Output_path)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    file_path = Output_path  # Dropbox에 저장될 경로와 파일명
    try:
        dbx.files_upload(doc_stream.getvalue(), file_path, mode=WriteMode('overwrite'))
        print(f"파일이 성공적으로 업로드되었습니다: {file_path}")
    except Exception as e:
        print(f"업로드 중 오류가 발생했습니다: {str(e)}")
    # 단일 파일 변환
    # 사용 예제

def CreateMemorizationBook(QuestionAnswer_json, AnswerCommentary_json, Output_path):
    # JSON 파일 열기
    with open(QuestionAnswer_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        QuestionAnswer_json_data = json.load(file)

    with open(AnswerCommentary_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        AnswerCommentary_json_data = json.load(file)

    dbx = dropbox.Dropbox(ACCESS_TOKEN)
    doc = Document()
    
    for i in range(len(QuestionAnswer_json_data)):
        if AnswerCommentary_json_data[i]['IsCorrect'] == False:
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(5)

            # 날짜, 월, 일
            now = datetime.now()
            current_date = now.date()
            cell = table.cell(0, 0)
            cell.text = "날짜"
            cell = table.cell(0, 1)
            cell.text = f"{current_date.year}년 {current_date.month}월 {current_date.day}일"
            
            # 과목
            cell = table.cell(1, 0)
            cell.text = "과목"
            cell = table.cell(1, 1)
            cell.text = "언어와 매체"

            cell = table.cell(2, 0)
            cell.text = "문제"
            cell = table.cell(2, 1)
            # 문제 + 지문 + 보기 5개 + 정답 + 해설
            cell.text = QuestionAnswer_json_data[i]['Question'] + "\n\n" + QuestionAnswer_json_data[i]['Text']

            cell = table.cell(3, 0)
            cell.text = "보기"
            cell = table.cell(3, 1)
            cell.text = "\n".join(QuestionAnswer_json_data[i]['Candidate'])
            # 난이도
            cell = table.cell(4, 0)
            cell.text = "정답"
            cell = table.cell(4, 1)
            cell.text = ""
            cell = table.cell(5, 0)
            cell.text = "해설"
            cell = table.cell(5, 1)
            cell.text = ""

            doc.add_page_break()
    
    #doc.save(Output_path)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    file_path = Output_path  # Dropbox에 저장될 경로와 파일명
    try:
        dbx.files_upload(doc_stream.getvalue(), file_path, mode=WriteMode('overwrite'))
        print(f"파일이 성공적으로 업로드되었습니다: {file_path}")
    except Exception as e:
        print(f"업로드 중 오류가 발생했습니다: {str(e)}")
    # 단일 파일 변환
    # 사용 예제

def load_pdfs_from_folder(folder_path):
    total_split_docs = []
    # 텍스트 분할기를 사용하여 문서를 분할합니다.
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

    for filename in os.listdir(folder_path):
        if filename.endswith('.pdf'):
            file_path = os.path.join(folder_path, filename)
            loader = PyPDFLoader(file_path)
            split_docs = loader.load_and_split(text_splitter=text_splitter)
            total_split_docs.extend(split_docs)
    return total_split_docs

def CreateCorrectAnswerNote(QuestionAnswer_json, AnswerCommentary_json, Output_path):
    with open(QuestionAnswer_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        QuestionAnswer_json_data = json.load(file)

    with open(AnswerCommentary_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        AnswerCommentary_json_data = json.load(file)

    folder_path = 'data/Insendium'
    pages = load_pdfs_from_folder(folder_path)
    dbx = dropbox.Dropbox(ACCESS_TOKEN)
    # VectorStore를 생성합니다.
    vector = FAISS.from_documents(pages, OpenAIEmbeddings())

    # Retriever를 생성합니다.
    retriever = vector.as_retriever()
    
    # hub에서 prompt를 가져옵니다 - 이 부분을 수정할 수 있습니다!
    prompt = hub.pull("hwchase17/openai-functions-agent")

    # Pydantic 모델 정의
    class AgentResponse(BaseModel):
        Question: str = Field(description="문제")
        Text: str = Field(description="문제에 첨부된 지문")
        Candidate: List[str] = Field(description="5가지 정답 후보")
        CorrectAnswer: int = Field(description="문제 정답")
        CommentarySummarize: List[str] = Field(description="문제 해설 3줄 요약")

    llm = ChatOpenAI(openai_api_key=openai_api_key2, model="gpt-4-turbo", temperature=0.1)
    structured_llm = llm.with_structured_output(AgentResponse)
    # parser가 get_format_instructions를 통해 스키마 생성, LLM 출력을 Json 형식으로 파싱
    parser = JsonOutputParser(pydantic_object=AgentResponse)

    doc = Document()
    
    for i in range(len(QuestionAnswer_json_data)):
        if AnswerCommentary_json_data[i]['IsCorrect'] == False:
            Question_Text = QuestionAnswer_json_data[i]['Question'] + QuestionAnswer_json_data[i]['Text']
            # langchain 패키지의 tools 모듈에서 retriever 도구를 생성
            retriever_tool = create_retriever_tool(
                retriever=retriever,
                name="pdf_search",
                # 도구에 대한 설명을 자세히 기입해야 합니다!!!
                description= Question_Text + "와 유사한 문서를 PDF 문서에서 검색합니다."
            )
            search = load_tools(["ddg-search"])[0]
            # tools 리스트에 search와 retriever_tool을 추가합니다.
            tools = [search, retriever_tool]
            
            # llm, tools, prompt를 인자로 사용합니다.
            agent = create_openai_functions_agent(structured_llm, tools, prompt)
            # AgentExecutor 클래스를 사용하여 agent와 tools를 설정하고, 상세한 로그를 출력하도록 verbose를 True로 설정합니다.
            agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)

            ########## 6. 채팅 기록을 수행하는 메모리를 추가합니다. ##########

            # 채팅 메시지 기록을 관리하는 객체를 생성합니다.
            message_history = ChatMessageHistory()

            # 채팅 메시지 기록이 추가된 에이전트를 생성합니다.
            agent_with_chat_history = RunnableWithMessageHistory(
                agent_executor,
                # 대부분의 실제 시나리오에서 세션 ID가 필요하기 때문에 이것이 필요합니다
                # 여기서는 간단한 메모리 내 ChatMessageHistory를 사용하기 때문에 실제로 사용되지 않습니다
                lambda session_id: message_history,
                # 프롬프트의 질문이 입력되는 key: "input"
                input_messages_key="input",
                output_messages_key="output",
                # 프롬프트의 메시지가 입력되는 key: "chat_history"
                history_messages_key="chat_history"
            )
            response = agent_with_chat_history.invoke(
                {
                    "input": Question_Text + """과 관련성이 높은 문제를 PDF 문서에서 찾고, 관련성이 높은 문제를 참고해서 유사한 문제를 만들어줘. 또한, 해설과 정답까지 알려줘.
                            Generate a JSON object with the following fields:
                            - Question: 문제
                            - Text: 문제에 첨부된 지문
                            - Candidate: 5가지 정답 후보
                            - CorrectAnswer: 실제 정답
                            - CommentarySummarize: 문제 해설 3줄 요약
                            
                            아래 예시처럼 작성해줘.
                            - Question: 밑줄 친 부분이 <보기>의 ㉠~㉤에 해당하는 예로 적절한 것은?
                            - Text: <보기>  
                    안은문장은 한 절이 그 속에 다른 절을 문장 성분의 하나로 안고 있는 문장이다. 이때 안겨 있는 절을 안긴절이라 하며, 안긴절의 종류에는 ㉠명사절, ㉡관형사절, ㉢부사절, ㉣서술절, ㉤인용절이 있다. 명사절, 관형사절, 부사절은 주로 전성 어미를 통해 실현된다. 서술절은 전성 어미 없이 실현되며, 인용절은 조사가 붙어 실현된다.
                            - Candidate: ["1.㉠: 그가 학교에 간다는 사실을 알고 있었다."
                                        "2.㉡: 내가 어제 본 영화는 매우 재미있었다.",
                                        "3.㉢: 비가 오기 전에 집에 도착해야 한다.",
                                        "4.㉣: 그녀가 도시에 도착했음을 확인했다.",
                                        "5.㉤: 그는 내일 회사를 그만둔다고 말했다."]
                            - CorrectAnswer: 5
                            - CommentarySummarize: 
                            ["토론의 내용은 아동의 개인 정보 노출과 관련된 법적 보호 필요성에 초점을 맞추고 있음.",
                            "찬성 측은 개인 정보 보호의 중요성을 강조하며 법적 보호를 주장할 수 있음.",
                            "반대 측은 개인 정보의 미숙한 권리 행사나 삭제 요구 권리의 제도화를 주장할 수 있음."]
                            """,
                    "output": parser
                },
                # 세션 ID를 설정합니다.
                # 여기서는 간단한 메모리 내 ChatMessageHistory를 사용하기 때문에 실제로 사용되지 않습니다
                config={"configurable": {"session_id": "MyTestSessionID"}}, 
            )
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(5)

            # 날짜, 월, 일
            now = datetime.now()
            current_date = now.date()
            cell = table.cell(0, 0)
            cell.text = "날짜"
            cell = table.cell(0, 1)
            cell.text = f"{current_date.year}년 {current_date.month}월 {current_date.day}일"
            
            # 과목
            cell = table.cell(1, 0)
            cell.text = "과목"
            cell = table.cell(1, 1)
            cell.text = "언어와 매체"

            cell = table.cell(2, 0)
            cell.text = "문제"
            cell = table.cell(2, 1)
            # 문제 + 지문 + 보기 5개 + 정답 + 해설
            #print(response['output'])
            print(response)
            resp = json.loads(response['output'])
            #cell.text = resp['Question'] + "\n\n" + resp['Text']
            cell.text = resp.Question + "\n\n" + resp.Text
            cell = table.cell(3, 0)
            cell.text = "보기"
            cell = table.cell(3, 1)
            cell.text = "\n".join(resp.Candidate)
            # 난이도
            cell = table.cell(4, 0)
            cell.text = "정답"
            cell = table.cell(4, 1)
            cell.text = str(resp.CorrectAnswer)
            cell = table.cell(5, 0)
            cell.text = "해설"
            cell = table.cell(5, 1)
            cell.text = "\n".join(resp.CommentarySummarize)

            doc.add_page_break()
    
    #doc.save(Output_path)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    file_path = Output_path  # Dropbox에 저장될 경로와 파일명
    try:
        dbx.files_upload(doc_stream.getvalue(), file_path, mode=WriteMode('overwrite'))
        print(f"파일이 성공적으로 업로드되었습니다: {file_path}")
    except Exception as e:
        print(f"업로드 중 오류가 발생했습니다: {str(e)}")

def MakeStudentInfoScoreCommentary(Question_pdf, Answer_json):
    response1 = MakeStudentInfo(Question_pdf = Question_pdf,
                Answer_json = Answer_json)
    
    response2 = MakeScoreCommentary(QuestionAnswerJson = response1)
    return response2

"""
def RAG_Chat(messages, 
             QuestionAnswer_json="./QuestionAnswer.json", 
             AnswerCommentary_json = "./AnswerCommentary.json"):

    cache_dir = LocalFileStore("./.cache/")

    splitter = CharacterTextSplitter.from_tiktoken_encoder(
        separator="\n",
        chunk_size=600,
        chunk_overlap=100,
    )
    loader = UnstructuredFileLoader("data/2024 수능특강 언어와 매체 정답.pdf")

    docs = loader.load_and_split(text_splitter=splitter)

    embeddings = OpenAIEmbeddings()

    cached_embeddings = CacheBackedEmbeddings.from_bytes_store(embeddings, cache_dir)

    vectorstore = FAISS.from_documents(docs, cached_embeddings)

    retriver = vectorstore.as_retriever()

    with open(QuestionAnswer_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        QuestionAnswer_json_data = json.load(file)

    with open(AnswerCommentary_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        AnswerCommentary_json_data = json.load(file)

    llm = ChatOpenAI(openai_api_key=openai_api_key, temperature=0.1)
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True, llm=llm) # BufferMemory 선언
    for i in range(len(QuestionAnswer_json_data)):
        if AnswerCommentary_json_data[i]['IsCorrect'] == False:

            memory.chat_memory.add_user_message(f"학생은 다음 문제를 틀렸다. 문제: {QuestionAnswer_json_data[i]['Question']}" 
                                                + "\n\n" + f"{QuestionAnswer_json_data[i]['Text']}"  
                                                + f"{QuestionAnswer_json_data[i]['Candidate']})"
                                                + f"학생이 제출한 정답: {AnswerCommentary_json_data[i]['StudentAnswer']}"
                                                + f"실제 정답: {AnswerCommentary_json_data[i]['CorrectAnswer']}") # 유저 메시지 추가
            # memory.chat_memory.add_ai_message("안녕하세요 무엇을 도와드릴까요?") # ai 메시지 추가
    prompt = ChatPromptTemplate.from_messages(
    [
        ("system", "You are a helpful chatbot"),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{input}"),
    ]
)
"""
    #memory.load_memory_variables({}) # 메모리 불러올 때, 빈 dict 넣기
    # {'history': 'Human: 안녕\nAI: 안녕하세요 무엇을 도와드릴까요?'}

    
"""
response1 = MakeStudentInfo(Question_pdf = "Epson-data/언어와매체10문제.pdf",
                Answer_json = "Epson-data/answers.json",
                Output_path = "Epson-data/QuestionAnswer.json")

response2 = MakeScoreCommentary(QuestionAnswerJson = "Epson-data/QuestionAnswer.json",
                               Output_path="Epson-data/AnswerCommentary.json")

#response = MakeStudentInfoScoreCommentary(Question_pdf = "Epson-data/언어와매체10문제.pdf",
#                               Answer_json = "Epson-data/answers.json")

CreateMemorizationBookAddCommentary(QuestionAnswer_json="Epson-data/QuestionAnswer.json",
                       AnswerCommentary_json="Epson-data/AnswerCommentary.json",
                       Output_path="/test/암기장.docx") # 원문제 + 정답 + 해설
      
CreateCorrectAnswerNote(QuestionAnswer_json="Epson-data/QuestionAnswer.json",
                       AnswerCommentary_json="Epson-data/AnswerCommentary.json",
                       Output_path="/test/유사문제.docx") # 유사문제 생성

CreateMemorizationBook(QuestionAnswer_json="Epson-data/QuestionAnswer.json",
                       AnswerCommentary_json="Epson-data/AnswerCommentary.json",
                       Output_path="/test/오답노트.docx") # 원문제만




# Initialize Dropbox client

def PostMakeStudentInfoScoreCommentary(data):
    dbx = dropbox.Dropbox(ACCESS_TOKEN)
    file_path = f"/test/{data['document']}"
    _, res = dbx.files_download(file_path)

    with open("temp.pdf", "wb") as f:
        f.write(res.content)
    doc = fitz.open("temp.pdf")
    
    #summary = finetuned_summarize(input_text.text)
    result = MakeStudentInfoScoreCommentary(Question_pdf=doc,
                                            Answer_json=data["submit"])
    return result





data = {
    "document": "언어와 매체 10문제.pdf",
    "submit": [
        {
            "question_num": 1,
            "answer": "1"
        },
        {
            "question_num": 2,
            "answer": "2"
        },
        {
            "question_num": 3,
            "answer": "3"
        },
        {
            "question_num": 4,
            "answer": "4"
        },
        {
            "question_num": 5,
            "answer": "5"
        },
        {
            "question_num": 6,
            "answer": "1"
        },
        {
            "question_num": 7,
            "answer": "2"
        },
        {
            "question_num": 8,
            "answer": "3"
        },
        {
            "question_num": 9,
            "answer": "4"
        },
        {
            "question_num": 10,
            "answer": "5"
        }
    ]
}

result = PostMakeStudentInfoScoreCommentary(data)
print(result)
"""